# Extracts raw text from a syllabus file so it can be handed to the LLM for
# schedule extraction. Supports PDF, Word (.docx), plain .txt, and screenshots
# (.jpg/.png, via OCR).
#
# SCREENSHOT/OCR SETUP: reading text out of an image needs the Tesseract OCR
# engine installed on your system (pytesseract is just a thin Python wrapper
# around it, and is already in requirements via pip):
#   Windows installer: https://github.com/UB-Mannheim/tesseract/wiki
#   After installing, either make sure tesseract.exe is on your PATH, or set
#   TESSERACT_CMD below to its full path (e.g.
#   r"C:\Program Files\Tesseract-OCR\tesseract.exe").
#
# extract_text_and_images() is the multimodal counterpart used by the Study
# Planner: slides often mix real text with diagrams/charts that OCR can't
# read at all, so it returns the images alongside the text (PDF pages
# rendered whole via PyMuPDF, embedded pictures pulled out of .docx files, or
# the file itself for a plain screenshot) so a vision-capable model can look
# at them directly rather than relying on OCR alone.

import base64
import io
from pathlib import Path

import fitz  # PyMuPDF
from pypdf import PdfReader
from docx import Document
from PIL import Image
import pytesseract

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # set via winget install, confirmed present
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png")

MAX_IMAGES_PER_FILE = 6
PDF_RENDER_ZOOM = 1.5  # ~150 DPI-equivalent -- legible for diagrams/charts without huge payloads


class SyllabusParseError(Exception):
    pass


def _sanitize_text(text):
    """Strips lone/unpaired UTF-16 surrogate codepoints (U+D800-U+DFFF)
    that some PDF fonts' broken/custom character maps can produce -- pypdf
    calling chr() on a raw code unit that was only ever half of a surrogate
    pair. Python's str type will silently hold these, but they're invalid
    Unicode: JSON-encoding one for an LM Studio request produces a
    technically-malformed escape that its parser correctly rejects,
    failing the whole request with an opaque server_error."""
    return "".join(ch for ch in text if not (0xD800 <= ord(ch) <= 0xDFFF))


def extract_text(path):
    """path: str or Path to a .pdf or .docx file. Returns the extracted text
    as a single string. Raises SyllabusParseError on unreadable/unsupported files."""
    path = Path(path)
    if not path.exists():
        raise SyllabusParseError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        text = _extract_pdf(path)
    elif ext == ".docx":
        text = _extract_docx(path)
    elif ext == ".txt":
        text = _extract_txt(path)
    elif ext in (".jpg", ".jpeg", ".png"):
        text = _extract_image(path)
    else:
        raise SyllabusParseError(
            f"Unsupported file type '{ext}' -- expected one of {SUPPORTED_EXTENSIONS}"
        )
    return _sanitize_text(text)


def _extract_pdf(path):
    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise SyllabusParseError(f"Couldn't open PDF {path.name}: {e}") from e

    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n".join(pages_text).strip()
    if not full_text:
        raise SyllabusParseError(
            f"No extractable text found in {path.name} -- it might be a scanned "
            "image rather than real text, which this doesn't handle."
        )
    return full_text


def _extract_docx(path):
    try:
        doc = Document(str(path))
    except Exception as e:
        raise SyllabusParseError(f"Couldn't open Word doc {path.name}: {e}") from e

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    # syllabi often put the actual date-bearing schedule in a table rather
    # than body paragraphs -- pull those in too, row by row
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n".join(parts).strip()
    if not full_text:
        raise SyllabusParseError(f"No extractable text found in {path.name}")
    return full_text


def _extract_txt(path):
    try:
        full_text = path.read_text(encoding="utf-8", errors="replace").strip()
    except Exception as e:
        raise SyllabusParseError(f"Couldn't read {path.name}: {e}") from e

    if not full_text:
        raise SyllabusParseError(f"No text found in {path.name}")
    return full_text


def _extract_image(path):
    try:
        image = Image.open(path)
        full_text = pytesseract.image_to_string(image).strip()
    except pytesseract.TesseractNotFoundError as e:
        raise SyllabusParseError(
            "Tesseract OCR isn't installed (or isn't on PATH). Install it from "
            "https://github.com/UB-Mannheim/tesseract/wiki, or set TESSERACT_CMD "
            "in syllabus_parser.py to its full path."
        ) from e
    except Exception as e:
        raise SyllabusParseError(f"Couldn't read {path.name}: {e}") from e

    if not full_text:
        raise SyllabusParseError(
            f"No readable text found in {path.name} -- try a clearer/higher-resolution screenshot."
        )
    return full_text


def _image_to_png_b64(image):
    buffer = io.BytesIO()
    image.convert("RGB").save(buffer, "PNG")
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def count_pdf_pages(path):
    """Total page count of a PDF -- e.g. for the Revision panel, which
    needs to know how many slides are IN an uploaded deck before "covered
    up to slide N" has anything real to validate/cap against."""
    path = Path(path)
    try:
        doc = fitz.open(str(path))
        try:
            return len(doc)
        finally:
            doc.close()
    except Exception as e:
        raise SyllabusParseError(f"Couldn't open PDF {path.name}: {e}") from e


def extract_text_and_images(path, max_images=MAX_IMAGES_PER_FILE, page_limit=None):
    """Like extract_text(), but for material that may need actual visual
    understanding (slides with diagrams/charts, not just OCR-able text).
    Returns (text, images_b64) where images_b64 is a list of base64-encoded
    PNG images meant to be handed to a vision-capable model alongside the
    text: whole-page renders for PDFs (covers vector diagrams OCR/text
    extraction would miss entirely), embedded pictures for Word docs, or the
    file itself for a plain screenshot.

    page_limit (PDF only): restrict extraction to just the first
    page_limit pages -- e.g. the Revision panel scoping this to "only the
    slides actually covered in class" rather than the whole deck. None
    (the default) means no restriction, same as before this param existed.

    Unlike extract_text(), this never raises just because a file has little
    or no extractable text -- an all-diagram slide deck is a normal case
    here, not an error."""
    path = Path(path)
    if not path.exists():
        raise SyllabusParseError(f"File not found: {path}")

    ext = path.suffix.lower()
    if ext == ".pdf":
        text, images = _extract_pdf_multimodal(path, max_images, page_limit=page_limit)
    elif ext == ".docx":
        text, images = _extract_docx_multimodal(path, max_images)
    elif ext == ".txt":
        text, images = _extract_txt(path), []
    elif ext in (".jpg", ".jpeg", ".png"):
        text, images = _extract_image_multimodal(path)
    else:
        raise SyllabusParseError(
            f"Unsupported file type '{ext}' -- expected one of {SUPPORTED_EXTENSIONS}"
        )
    return _sanitize_text(text), images


MIN_DRAWINGS_FOR_VISUAL_PAGE = 8  # a plain text slide's occasional background rectangle or bullet dash
                                   # won't reach this; a real diagram/flowchart/table drawn out of vector
                                   # shapes usually will -- see _page_has_visual_content()


def _page_has_visual_content(page):
    """True if a PDF page likely has a chart/diagram/photo actually worth
    LOOKING at, as opposed to a plain text slide whose content is already
    fully captured by extract_text() -- an embedded raster image (a photo,
    a chart exported as a PNG/JPEG, a scanned figure) is a reliable signal
    on its own; a page built almost entirely from vector shapes (a
    flowchart/diagram drawn with lines and boxes rather than an embedded
    picture) is inferred from an unusually high drawing-path count."""
    if page.get_images(full=True):
        return True
    return len(page.get_drawings()) >= MIN_DRAWINGS_FOR_VISUAL_PAGE


def _extract_pdf_multimodal(path, max_images, page_limit=None):
    try:
        reader = PdfReader(str(path))
        pages = reader.pages[:page_limit] if page_limit else reader.pages
        text = "\n".join((page.extract_text() or "") for page in pages).strip()
    except Exception as e:
        raise SyllabusParseError(f"Couldn't open PDF {path.name}: {e}") from e

    images_b64 = []
    try:
        doc = fitz.open(str(path))
        try:
            total_pages = min(len(doc), page_limit) if page_limit else len(doc)
            # Only render pages that actually have a chart/diagram/photo --
            # a plain text slide's content already lives in `text` above,
            # so re-sending it as a whole-page image just burns vision
            # tokens (and, on modest local hardware, real wall-clock time --
            # see llm_client._chat_multimodal()'s halving/timeout retry)
            # for nothing the model doesn't already have.
            visual_indices = [i for i in range(total_pages) if _page_has_visual_content(doc[i])]

            if len(visual_indices) > max_images > 1:
                # More visually-rich pages than we can afford to send --
                # sample evenly across THEM specifically (not the whole
                # deck), so a cluster of diagrams late in the covered
                # range isn't starved just because earlier pages happened
                # to be text-only.
                last = len(visual_indices) - 1
                indices = [visual_indices[round(i * last / (max_images - 1))] for i in range(max_images)]
            else:
                indices = visual_indices[:max_images]

            for page_index in indices:
                pixmap = doc[page_index].get_pixmap(matrix=fitz.Matrix(PDF_RENDER_ZOOM, PDF_RENDER_ZOOM))
                images_b64.append(base64.b64encode(pixmap.tobytes("png")).decode("ascii"))
        finally:
            doc.close()
    except Exception:
        pass  # rendering is best-effort -- fall back to whatever text was extracted

    return text, images_b64


def _extract_docx_multimodal(path, max_images):
    try:
        doc = Document(str(path))
    except Exception as e:
        raise SyllabusParseError(f"Couldn't open Word doc {path.name}: {e}") from e

    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()

    images_b64 = []
    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            images_b64.append(_image_to_png_b64(Image.open(io.BytesIO(rel.target_part.blob))))
        except Exception:
            continue
        if len(images_b64) >= max_images:
            break

    return text, images_b64


def _extract_image_multimodal(path):
    try:
        image = Image.open(path)
    except Exception as e:
        raise SyllabusParseError(f"Couldn't read {path.name}: {e}") from e

    try:
        text = pytesseract.image_to_string(image).strip()
    except pytesseract.TesseractNotFoundError:
        text = ""  # OCR unavailable -- the raw image can still go to a vision model
    except Exception:
        text = ""

    return text, [_image_to_png_b64(image)]
